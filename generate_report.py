"""按学术论文模板格式生成项目报告（大作业参考格式模板.doc）"""
import os, copy
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RESULTS_DIR = r"D:\NNU\语音识别技术\大作业\results"
OUT_PATH    = r"D:\NNU\语音识别技术\大作业\项目报告_基于声音的鸟类识别系统_v3.docx"

doc = Document()

# ══════════════════════════════════════════════════════════════
# 页面设置（参照模板）
# ══════════════════════════════════════════════════════════════
sec = doc.sections[0]
sec.top_margin    = Cm(0.5)
sec.bottom_margin = Cm(0.5)
sec.left_margin   = Cm(1.6)
sec.right_margin  = Cm(1.6)
sec.page_width    = Cm(20.14)
sec.page_height   = Cm(27.55)

# 双栏排版
def set_two_columns(section):
    sectPr = section._sectPr
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '425')  # ~0.75cm 间距
    sectPr.append(cols)

# 正文双栏设置在引言前，先单栏写标题/摘要
# ══════════════════════════════════════════════════════════════
# 辅助函数
# ══════════════════════════════════════════════════════════════
def para(text="", font="宋体", size=10.5, bold=False,
         align=WD_ALIGN_PARAGRAPH.LEFT, color=None, space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.font.name = font
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = color
        _set_eastasia(run, font)
    return p

def mixed_para(parts, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=0):
    """parts = [(text, font, size, bold), ...]"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    for (text, font, size, bold) in parts:
        run = p.add_run(text)
        run.font.name = font
        run.font.size = Pt(size)
        run.bold = bold
        _set_eastasia(run, font)
    return p

def _set_eastasia(run, font_name):
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)

def heading1(text):
    """一级标题：小四黑体（12pt）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.name = "黑体"; run.font.size = Pt(12); run.bold = True
    _set_eastasia(run, "黑体")
    return p

def heading2(text):
    """二级标题：五号黑体（10.5pt）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    run.font.name = "黑体"; run.font.size = Pt(10.5); run.bold = True
    _set_eastasia(run, "黑体")
    return p

def body(text, indent=True):
    """正文：五号宋体（10.5pt），首行缩进2字"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Pt(21)   # 约2个字
    run = p.add_run(text)
    run.font.name = "宋体"; run.font.size = Pt(10.5)
    _set_eastasia(run, "宋体")
    return p

def fig_caption(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = "宋体"
    p.runs[0].font.size = Pt(9)
    _set_eastasia(p.runs[0], "宋体")
    return p

def table_caption(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = "宋体"
    p.runs[0].font.size = Pt(9)
    p.runs[0].bold = True
    _set_eastasia(p.runs[0], "宋体")
    return p

def add_table(headers, rows, font_size=9):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Table Grid"
    # header row
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.paragraphs[0].clear()
        run = c.paragraphs[0].add_run(h)
        run.font.name = "黑体"; run.font.size = Pt(font_size); run.bold = True
        _set_eastasia(run, "黑体")
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            c.paragraphs[0].clear()
            run = c.paragraphs[0].add_run(str(val))
            run.font.name = "宋体"; run.font.size = Pt(font_size)
            _set_eastasia(run, "宋体")
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return t

def add_figure(path, width=Cm(8.5)):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=width)

def section_break_to_two_col():
    """插入连续分节符，之后设为双栏"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    sectPr = OxmlElement('w:sectPr')
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '425')
    sectPr.append(cols)
    type_elem = OxmlElement('w:type')
    type_elem.set(qn('w:val'), 'continuous')
    sectPr.insert(0, type_elem)
    pPr.append(sectPr)

def section_break_to_one_col():
    """插入连续分节符，切回单栏（用于参考文献）"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    sectPr = OxmlElement('w:sectPr')
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '1')
    sectPr.append(cols)
    type_elem = OxmlElement('w:type')
    type_elem.set(qn('w:val'), 'continuous')
    sectPr.insert(0, type_elem)
    pPr.append(sectPr)

# ══════════════════════════════════════════════════════════════
# ① 文章编号行
# ══════════════════════════════════════════════════════════════
para("文章编号：（留空）", font="Arial", size=9.5)

# ══════════════════════════════════════════════════════════════
# ② 中文标题
# ══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(4)
run = p.add_run("基于声音的鸟类识别系统")
run.font.name = "黑体"; run.font.size = Pt(16); run.bold = True
_set_eastasia(run, "黑体")

# ══════════════════════════════════════════════════════════════
# ③ 作者
# ══════════════════════════════════════════════════════════════
para("作者姓名", font="仿宋_GB2312", size=14,
     align=WD_ALIGN_PARAGRAPH.CENTER)

# ④ 单位
para("（南京师范大学  学院，江苏 南京  邮编）",
     font="宋体", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# ⑤ 摘要（楷体小五号）
# ══════════════════════════════════════════════════════════════
mixed_para([
    ("摘要：", "黑体", 9, True),
    ("本文基于BirdCLEF 2026数据集，构建了一套完整的基于声音的鸟类识别系统。"
     "首先对音频进行预处理（重采样至22050 Hz、截取/填充至5秒），"
     "提取包括梅尔频率倒谱系数（MFCC）及其一阶、二阶差分、音高（F0）、"
     "音量（RMS）、音色（频谱质心/带宽/滚降/平坦度）和过零率（ZCR）"
     "等共278维声学特征。在此基础上，分别构建了K近邻（KNN）、"
     "支持向量机（SVM）、随机森林等传统机器学习模型，"
     "以及基于对数Mel频谱图的卷积神经网络（CNN）"
     "和基于MFCC序列的双向长短期记忆网络（BiLSTM）深度学习模型。"
     "实验结果表明，CNN模型取得最优性能，测试集准确率为48.75%，"
     "Macro F1分数为0.463，优于其他所有模型。"
     "最后，本文使用Gradio构建了Web交互界面，"
     "支持实时音频上传与多模型识别结果展示。",
     "楷体", 9, False),
])

mixed_para([
    ("关键词：", "黑体", 9, True),
    ("鸟类识别；声音识别；梅尔频率倒谱系数；卷积神经网络；双向长短期记忆网络",
     "楷体", 9, False),
])

mixed_para([
    ("中图分类号：TP391     ", "楷体", 9, True),
    ("文献标识码：A", "楷体", 9, True),
])

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# ⑥ 英文标题
# ══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Bird Species Recognition System Based on Sound")
run.font.name = "Times New Roman"; run.font.size = Pt(14); run.bold = True

para("Author Name", font="Times New Roman", size=12,
     align=WD_ALIGN_PARAGRAPH.CENTER)
para("(School Name, Nanjing Normal University, Nanjing 210097, China)",
     font="Times New Roman", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

mixed_para([
    ("Abstract: ", "Times New Roman", 9, True),
    ("This paper presents a bird sound recognition system based on the BirdCLEF 2026 dataset. "
     "Audio clips are preprocessed (resampled to 22050 Hz, padded/trimmed to 5 seconds) and "
     "278-dimensional acoustic features are extracted, including MFCCs with first- and "
     "second-order deltas, pitch (F0), volume (RMS energy), timbre (spectral centroid, "
     "bandwidth, rolloff, flatness), and zero-crossing rate (ZCR). Three traditional machine "
     "learning models (KNN, SVM, Random Forest) and two deep learning models (CNN on "
     "log-Mel spectrograms and BiLSTM on MFCC sequences) are built and compared. "
     "Experimental results show that CNN achieves the best performance with 48.75% test "
     "accuracy and a Macro F1 score of 0.463. A Gradio-based web interface is developed "
     "for real-time bird sound recognition.",
     "Times New Roman", 9, False),
])
mixed_para([
    ("Key words: ", "Times New Roman", 9, True),
    ("bird recognition; sound recognition; MFCC; CNN; BiLSTM",
     "Times New Roman", 9, False),
])

# ══════════════════════════════════════════════════════════════
# 切换双栏
# ══════════════════════════════════════════════════════════════
section_break_to_two_col()

# ══════════════════════════════════════════════════════════════
# 0  引言
# ══════════════════════════════════════════════════════════════
heading1("0  引言")
body("鸟类是生态系统健康的重要指示物种，其多样性监测对于生物保护具有重要意义。"
     "传统人工野外调查效率低、成本高，难以大规模部署。"
     "近年来，被动声学监测技术的兴起使得自动化鸟鸣识别成为可能[1]。")
body("BirdCLEF 2026是Kaggle举办的国际鸟类声音识别竞赛，"
     "提供了来自南美洲热带地区的大规模野外鸟鸣录音数据集。"
     "本文基于该数据集，综合运用传统机器学习与深度学习方法，"
     "构建多模型对比系统并提供Web交互界面，"
     "旨在探索不同特征表示与模型结构对鸟鸣识别性能的影响。")

# ══════════════════════════════════════════════════════════════
# 1  数据集与预处理
# ══════════════════════════════════════════════════════════════
heading1("1  数据集与预处理")

heading2("1.1  数据集")
body("本文使用BirdCLEF 2026数据集，从中筛选鸟类（Aves）样本，"
     "选取样本数量最多的20个物种进行实验。"
     "每类深度学习最多取100条、传统机器学习最多取50条，"
     "共约1600条训练样本，音频格式为.ogg。")

heading2("1.2  预处理")
body("（1）重采样：所有音频统一重采样至22050 Hz，消除采样率差异；")
body("（2）单声道：立体声混合为单声道；")
body("（3）时长标准化：截取或零填充至固定5秒（110250采样点）；")
body("（4）数据划分：按8:1:1比例进行分层随机划分（训练/验证/测试集），"
     "保证各子集类别比例一致。")

add_figure(os.path.join(RESULTS_DIR, "03_audio_analysis.png"), Cm(8.5))
fig_caption("图1  鸟鸣音频波形、Mel频谱、MFCC与色度图")

# ══════════════════════════════════════════════════════════════
# 2  特征提取
# ══════════════════════════════════════════════════════════════
heading1("2  特征提取")
body("本文针对不同模型需求，提取两类特征表示。")

heading2("2.1  传统ML统计特征向量（278维）")
body("对每条5秒音频计算以下声学特征在时间轴上的均值与标准差，"
     "拼接成固定长度的特征向量。")

table_caption("表1  传统机器学习特征构成（278维）")
add_table(
    ["特征类别", "描述", "维度"],
    [
        ["MFCC+Δ+ΔΔ", "40阶MFCC及一、二阶差分（均值+标准差）", "240"],
        ["Chroma",    "12维色度特征，反映音调内容", "24"],
        ["频谱形状",   "质心、带宽、滚降点、平坦度", "8"],
        ["ZCR",       "过零率，反映鸣叫速率（rate）", "2"],
        ["RMS",       "均方根能量，反映音量（volume）", "2"],
        ["F0 (YIN)",  "基频，反映音高（pitch），32–2048 Hz", "2"],
        ["合计",      "—", "278"],
    ]
)
doc.add_paragraph()
body("其中，音高（F0）使用YIN算法估计，"
     "仅统计有声帧（f0>0）的均值与标准差，"
     "静音帧置零。音色（timbre）由频谱质心、带宽、"
     "滚降点和平坦度共同刻画[2]。")

heading2("2.2  CNN输入：Log-Mel频谱图")
body("参数设置n_mels=128、n_fft=2048、hop_length=512，"
     "计算对数梅尔频谱图并沿时间轴固定为212帧，"
     "归一化至[0,1]，形状为(1, 128, 212)。")

heading2("2.3  BiLSTM输入：MFCC时序矩阵")
body("提取40阶MFCC时序序列，对每个倒谱系数"
     "独立进行z-score标准化，形状为(212, 40)，"
     "保留了鸣声随时间变化的动态信息[3]。")

add_figure(os.path.join(RESULTS_DIR, "05_pitch_energy_zcr.png"), Cm(8.5))
fig_caption("图2  基频（Pitch）、音量（RMS）与过零率时序曲线")

# ══════════════════════════════════════════════════════════════
# 3  模型构建
# ══════════════════════════════════════════════════════════════
heading1("3  模型构建")

heading2("3.1  传统机器学习模型")
body("三种传统ML模型均以278维特征向量为输入，"
     "训练前经StandardScaler标准化，"
     "采用5折分层交叉验证（StratifiedKFold）"
     "结合GridSearchCV进行超参数搜索。")
body("KNN通过欧氏距离寻找K个最近邻居进行多数投票，"
     "网格搜索K∈{3,5,7,9}及距离权重策略。")
body("SVM使用RBF核将特征映射到高维空间，"
     "寻找最大间隔决策超平面，"
     "网格搜索C∈{0.1,1,10}和γ∈{scale,auto}[4]。")
body("随机森林集成多棵决策树，"
     "每棵树用Bootstrap样本和随机子特征训练，"
     "最终多数投票，有效缓解过拟合[5]。")

heading2("3.2  卷积神经网络（CNN）")
body("以Log-Mel频谱图作为二维图像输入，"
     "模拟图像分类思路提取时频局部特征。"
     "网络由4个卷积块构成（Conv2d-BN-ReLU×2+MaxPool+Dropout），"
     "通道数依次为32→64→128→256，"
     "之后接AdaptiveAvgPool2d(4×4)和全连接分类头（4096→512→256→20）。"
     "采用AdamW优化器（lr=1e-3，weight_decay=1e-4）"
     "和CosineAnnealingLR学习率调度，"
     "使用WeightedRandomSampler处理类别不平衡，"
     "以验证集准确率为指标进行早停（patience=7）[6]。")

heading2("3.3  双向LSTM + 注意力机制（BiLSTM）")
body("将MFCC时序矩阵作为序列输入，"
     "通过双向LSTM同时建模正向与反向时序依赖。"
     "引入注意力机制对各时刻隐状态加权聚合，"
     "突出关键鸣叫片段（hidden_size=128，2层，dropout=0.3），"
     "全连接头256→128→20[3]。训练策略与CNN相同。")

# ══════════════════════════════════════════════════════════════
# 4  实验结果与分析
# ══════════════════════════════════════════════════════════════
heading1("4  实验结果与分析")

heading2("4.1  传统机器学习结果")
body("表2给出了三种传统ML模型在5折交叉验证"
     "和独立测试集上的性能。SVM取得最优结果，"
     "测试准确率36%，Macro F1为0.355。"
     "KNN性能最差（22%），主要原因是"
     "278维高维特征空间中距离集中现象降低了近邻搜索的区分度。")

table_caption("表2  传统机器学习模型性能对比")
add_table(
    ["模型", "CV均值", "CV标准差", "测试准确率", "Macro F1"],
    [
        ["KNN",   "18.75%", "±1.48%", "22.00%", "0.209"],
        ["SVM",   "33.13%", "±1.63%", "36.00%", "0.355"],
        ["RF",    "29.88%", "±3.15%", "32.50%", "0.314"],
    ]
)
doc.add_paragraph()

add_figure(os.path.join(RESULTS_DIR, "cm_svm.png"), Cm(8.0))
fig_caption("图3  SVM混淆矩阵（测试集，20类）")

heading2("4.2  深度学习结果")
body("表3为两种深度学习模型的性能。"
     "CNN以48.75%的测试准确率成为所有模型中最优者，"
     "验证了二维卷积网络在处理频谱图像方面的优势。"
     "BiLSTM以38.75%紧随其后，优于全部传统ML方法。")

table_caption("表3  深度学习模型性能对比")
add_table(
    ["模型", "测试准确率", "Macro F1", "输入"],
    [
        ["CNN",    "48.75%", "0.463", "Log-Mel (1×128×212)"],
        ["BiLSTM", "38.75%", "0.385", "MFCC (212×40)"],
    ]
)
doc.add_paragraph()

add_figure(os.path.join(RESULTS_DIR, "cnn_training_curves.png"), Cm(8.5))
fig_caption("图4  CNN训练曲线（损失与准确率随Epoch变化）")

add_figure(os.path.join(RESULTS_DIR, "cnn_confusion_matrix.png"), Cm(8.0))
fig_caption("图5  CNN混淆矩阵（测试集，20类）")

heading2("4.3  综合对比与分析")

table_caption("表4  五种模型综合性能排名")
add_table(
    ["排名", "模型", "类型", "测试准确率", "Macro F1"],
    [
        ["1", "CNN",   "深度学习", "48.75%", "0.463"],
        ["2", "BiLSTM","深度学习","38.75%", "0.385"],
        ["3", "SVM",   "传统ML",  "36.00%", "0.355"],
        ["4", "RF",    "传统ML",  "32.50%", "0.314"],
        ["5", "KNN",   "传统ML",  "22.00%", "0.209"],
    ]
)
doc.add_paragraph()

add_figure(os.path.join(RESULTS_DIR, "final_comparison.png"), Cm(8.5))
fig_caption("图6  五种模型测试准确率与Macro F1综合对比")

body("深度学习整体优于传统ML，"
     "主要原因是CNN和BiLSTM能够直接从原始时频表示中"
     "自动学习判别特征，而传统ML依赖手工统计特征，"
     "丢失了大量时序结构信息。"
     "CNN优于BiLSTM，是因为Log-Mel频谱图"
     "保留了完整的二维时频结构，"
     "卷积核可高效捕获局部时频模式。")
body("整体准确率偏低的原因主要有：①每类训练样本仅50–100条，"
     "深度模型难以充分拟合；②部分物种鸣声频率重叠，"
     "区分难度较高（混淆矩阵中Osprey与Tropical Screech Owl"
     "之间存在明显混淆）；③实验在CPU上进行，"
     "模型规模和训练轮数受限。")

# ══════════════════════════════════════════════════════════════
# 5  用户界面
# ══════════════════════════════════════════════════════════════
heading1("5  用户界面")
body("本文使用Gradio 4.x构建了基于Web的交互式识别界面，"
     "支持上传.ogg/.wav/.mp3格式鸟鸣录音，"
     "一键运行全部5个模型并同时展示结果。"
     "界面主要包含：音频上传与播放区、"
     "音频可视化面板（波形/Log-Mel频谱/MFCC三合一图）、"
     "声学特征摘要（时长/平均音高/音量/过零率）、"
     "以及各模型Top-5物种识别结果与置信度。"
     "运行方式：python 06_ui.py，访问http://localhost:7860。")

# ══════════════════════════════════════════════════════════════
# 6  结论
# ══════════════════════════════════════════════════════════════
heading1("6  结论")
body("本文基于BirdCLEF 2026数据集实现了完整的鸟类鸣声识别系统，"
     "涵盖数据预处理、多维特征提取、"
     "5种模型（KNN/SVM/RF/CNN/BiLSTM）的构建与对比"
     "及Web交互界面。实验表明CNN（48.75%）取得最优性能，"
     "深度学习整体优于传统机器学习。"
     "未来可通过数据增强、迁移学习（如BirdNET预训练模型）"
     "和多模型集成进一步提升识别准确率，"
     "并引入Transformer等更强序列建模结构。")

# ══════════════════════════════════════════════════════════════
# 切回单栏（参考文献）
# ══════════════════════════════════════════════════════════════
section_break_to_one_col()

# ══════════════════════════════════════════════════════════════
# 参考文献
# ══════════════════════════════════════════════════════════════
heading1("参考文献")
refs = [
    "[1] Stowell D, Wood M D, Pamuła H, et al. Automatic acoustic detection of birds through deep learning: the first Bird Audio Detection challenge[J]. Methods in Ecology and Evolution, 2019, 10(3): 368-380.",
    "[2] McFee B, Raffel C, Liang D, et al. librosa: Audio and music signal analysis in python[C]//Proceedings of the 14th Python in Science Conference. 2015: 18-25.",
    "[3] Hochreiter S, Schmidhuber J. Long short-term memory[J]. Neural Computation, 1997, 9(8): 1735-1780.",
    "[4] Cortes C, Vapnik V. Support-vector networks[J]. Machine Learning, 1995, 20(3): 273-297.",
    "[5] Breiman L. Random forests[J]. Machine Learning, 2001, 45(1): 5-32.",
    "[6] He K, Zhang X, Ren S, et al. Deep residual learning for image recognition[C]//Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition. 2016: 770-778.",
    "[7] Kahl S, Wood C M, Eibl M, et al. BirdNET: A deep learning solution for avian diversity monitoring[J]. Ecological Informatics, 2021, 61: 101236.",
]
for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.first_line_indent = Pt(0)
    p.runs[0].font.name = "宋体"
    p.runs[0].font.size = Pt(9)
    _set_eastasia(p.runs[0], "宋体")

# ══════════════════════════════════════════════════════════════
# 作者简介
# ══════════════════════════════════════════════════════════════
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("作者简介：")
run.font.name = "黑体"; run.font.size = Pt(7.5); run.bold = True
_set_eastasia(run, "黑体")
run2 = p.add_run("姓名（出生年—），学历，主要研究领域为语音识别与机器学习。")
run2.font.name = "宋体"; run2.font.size = Pt(7.5)
_set_eastasia(run2, "宋体")

p2 = doc.add_paragraph("E-mail：")
p2.runs[0].font.name = "宋体"; p2.runs[0].font.size = Pt(7.5)
_set_eastasia(p2.runs[0], "宋体")

# ══════════════════════════════════════════════════════════════
# 保存
# ══════════════════════════════════════════════════════════════
doc.save(OUT_PATH)
print(f"报告已生成：{OUT_PATH}")
